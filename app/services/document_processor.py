import os
import io
import hashlib
import json
from datetime import datetime
from typing import BinaryIO, Dict, Any, Optional, Tuple
from fastapi import UploadFile, HTTPException
import networkx as nx
import matplotlib.pyplot as plt
from PIL import Image
import pytesseract
import pydot
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PyPDF2 import PdfReader
from .storage.base import BaseStorage, StorageType
from .storage.local import LocalStorage
from .storage.s3 import S3Storage
from .database import DatabaseManager

class DocumentProcessor:
    def __init__(
        self,
        storage_type: StorageType = StorageType.LOCAL,
        storage_config: Optional[Dict[str, Any]] = None,
        db_path: str = 'sqlite:///./data/documents.db'
    ):
        self.storage_config = storage_config or {}
        
        # Initialize storage backend
        if storage_type == StorageType.LOCAL:
            base_path = storage_config.get('base_path', './data/uploads')
            self.storage = LocalStorage(base_path)
        else:  # S3
            self.storage = S3Storage(
                bucket_name=storage_config['bucket_name'],
                aws_access_key_id=storage_config.get('aws_access_key_id'),
                aws_secret_access_key=storage_config.get('aws_secret_access_key'),
                endpoint_url=storage_config.get('endpoint_url'),
                region_name=storage_config.get('region_name')
            )
        
        # Initialize database
        self.db = DatabaseManager(db_path)
    
    def _generate_file_path(self, original_filename: str, session_id: str, doc_type: str) -> str:
        """Generate a unique file path for storage"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{original_filename}"
        return os.path.join(session_id, doc_type, filename)
    
    def _extract_content_by_format(self, content: bytes, content_type: str) -> str:
        """Extract text content based on file format"""
        format_handlers = {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_xlsx,
            'application/pdf': self._extract_pdf
        }
        
        if content_type not in format_handlers:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {content_type}. Only DOCX, XLSX, and PDF files are supported."
            )
        
        return format_handlers[content_type](content)
    
    def _extract_docx(self, content: bytes) -> str:
        """Extract text content from DOCX file"""
        doc = DocxDocument(io.BytesIO(content))
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_xlsx(self, content: bytes) -> str:
        """Extract text content from XLSX file"""
        wb = load_workbook(io.BytesIO(content))
        texts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows():
                row_text = ' '.join(str(cell.value) if cell.value is not None else '' for cell in row)
                if row_text.strip():
                    texts.append(row_text)
        return '\n'.join(texts)

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text content from PDF file"""
        pdf = PdfReader(io.BytesIO(content))
        texts = [page.extract_text() for page in pdf.pages]
        return '\n'.join(texts)

    def _extract_images_from_content(self, content: bytes, content_type: str) -> List[Image.Image]:
        """Extract images from document content based on file type"""
        images = []
        try:
            if content_type == 'application/pdf':
                pdf = PdfReader(io.BytesIO(content))
                for page in pdf.pages:
                    if '/XObject' in page['/Resources']:
                        for obj in page['/Resources']['/XObject'].get_object().values():
                            if obj['/Subtype'] == '/Image':
                                img_data = obj.get_object().get_data()
                                img = Image.open(io.BytesIO(img_data))
                                images.append(img)
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                doc = DocxDocument(io.BytesIO(content))
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        img_data = rel.target_part.blob
                        img = Image.open(io.BytesIO(img_data))
                        images.append(img)
        except Exception as e:
            print(f"Error extracting images: {str(e)}")
        return images

    def _detect_diagram_type(self, image: Image.Image) -> Optional[str]:
        """Detect the type of technical diagram in the image"""
        # TODO: Implement ML-based diagram type detection
        # For now, use basic image analysis to guess the type
        # This should be replaced with a proper ML model
        
        # Convert image to grayscale for analysis
        gray = image.convert('L')
        # Basic analysis based on image characteristics
        # This is a placeholder and should be replaced with proper ML detection
        return 'network_topology'  # Placeholder return

    def _process_technical_diagram(self, image: Image.Image, diagram_type: str) -> Dict[str, str]:
        """Convert technical diagram to descriptive text based on type"""
        prompts = {
            'network_topology': "Describe this network topology diagram including all nodes, connections, and their relationships. Include details about network segments, devices, and protocols if visible.",
            'class_diagram': "Describe this UML class diagram including all classes, their attributes, methods, and relationships. Include inheritance, composition, and other connections between classes.",
            'sequence_diagram': "Describe this sequence diagram including all actors, systems, and the sequence of interactions between them. Include the flow of messages and their timing.",
            'data_schema': "Describe this data schema including all entities, their attributes, relationships, and cardinality. Include primary keys, foreign keys, and any constraints visible."
        }
        
        if diagram_type not in prompts:
            return {}
            
        # TODO: Implement diagram-to-text conversion using computer vision and LLM
        # For now, return a placeholder
        return {
            'type': diagram_type,
            'description': f"Placeholder description for {diagram_type}"
        }

    def _process_document_content(self, content: bytes, content_type: str, doc_type: str) -> Dict[str, Any]:
        """Process document content based on document type and extract technical diagrams"""
        processed_data = {}
        
        # Extract images from the document
        images = self._extract_images_from_content(content, content_type)
        
        # Process each image for technical diagrams
        technical_diagrams = []
        for img in images:
            diagram_type = self._detect_diagram_type(img)
            if diagram_type:
                diagram_info = self._process_technical_diagram(img, diagram_type)
                if diagram_info:
                    technical_diagrams.append(diagram_info)
        
        if technical_diagrams:
            processed_data['technical_diagrams'] = technical_diagrams
        
        return processed_data

    def _prepare_metadata(self, file: UploadFile, session_id: str, doc_type: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Prepare metadata for document storage"""
        base_metadata = metadata or {}
        base_metadata.update({
            'original_filename': file.filename,
            'session_id': session_id,
            'doc_type': doc_type,
            'upload_time': datetime.now().isoformat()
        })
        return base_metadata

    async def save_document(
        self,
        file: UploadFile,
        session_id: str,
        doc_type: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """Save a document, hash it, extract content, and store in database"""
        file_path = self._generate_file_path(file.filename, session_id, doc_type)
        
        # Read file content and compute hash
        content = await file.read()
        file_hash = hashlib.sha256(content).hexdigest()
        
        # Map content types to extraction functions
        content_extractors = {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_content,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_xlsx_content,
            'application/pdf': self._extract_pdf_content
        }
        
        # Extract content
        extracted_content = None
        content_type = file.content_type.lower()
        
        if content_type not in content_extractors:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {content_type}. Only DOCX, XLSX, and PDF files are supported."
            )
            
        try:
            extracted_content = content_extractors[content_type](content)
        except Exception as e:
            print(f"Content extraction failed: {str(e)}")
        
        # Save to database
        doc_data = {
            'id': file_path,
            'file_hash': file_hash,
            'content': extracted_content,
            'original_filename': file.filename,
            'session_id': session_id,
            'doc_type': doc_type,
            'metadata': json.dumps(metadata) if metadata else None
        }
        self.db.save_document(doc_data)
        
        # Reset file position for storage
        await file.seek(0)
        
        # Prepare metadata
        metadata = self._prepare_metadata(file, session_id, doc_type, metadata)
        
        # Process specific document types
        if doc_type == "network_topology":
            graph_data = self._process_network_topology(io.BytesIO(content))
            metadata['graph_data'] = graph_data
        
        # Save file
        stored_path = self.storage.save_file(io.BytesIO(content), file_path, metadata)
        
        return stored_path, metadata
    
    def _process_network_topology(self, file_obj: BinaryIO) -> Dict[str, Any]:
        """Process network topology image and extract graph data"""
        # Create a temporary file for image processing
        img = Image.open(file_obj)
        
        # Extract text from image using OCR
        text = pytesseract.image_to_string(img)
        
        # Create graph from image
        graph = nx.Graph()
        
        # Parse the text to identify nodes and edges
        # This is a simple example - you might need more sophisticated parsing
        lines = text.split('\n')
        for line in lines:
            if '->' in line:
                source, target = line.split('->')
                graph.add_edge(source.strip(), target.strip())
            elif node := line.strip():
                graph.add_node(node)
        
        # Convert graph to dot format
        dot_data = nx.drawing.nx_pydot.to_pydot(graph)
        
        return {
            'nodes': list(graph.nodes()),
            'edges': list(graph.edges()),
            'dot_data': dot_data.to_string()
        }
    
    def get_document_url(self, file_path: str, expires_in: int = 3600) -> str:
        """Get URL for accessing the document"""
        return self.storage.get_file_url(file_path, expires_in)
    
    def get_document(self, file_path: str) -> BinaryIO:
        """Get document content"""
        return self.storage.get_file(file_path)
    
    def delete_document(self, file_path: str) -> None:
        """Delete a document"""
        self.storage.delete_file(file_path)
